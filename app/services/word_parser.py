"""Word文档解析器"""
import re
from typing import List, Optional
from docx import Document
from docx.table import Table
from app.models import GuestInfo


class WordParser:
    """Word文档解析器 - 支持.docx格式"""
    
    def __init__(self):
        self.guests: List[GuestInfo] = []
        self.text_content: str = ''
    
    def parse_docx(self, file_path: str) -> tuple[str, List[GuestInfo]]:
        """
        解析.docx文件
        返回: (文本内容, 来宾列表)
        """
        try:
            doc = Document(file_path)
            
            # 提取所有文本
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            self.text_content = '\n'.join(text_parts)
            
            # 提取表格中的来宾信息
            self.guests = []
            for table in doc.tables:
                self._extract_guests_from_table(table)
            
            return self.text_content, self.guests
            
        except Exception as e:
            print(f'解析Word文件失败: {e}')
            return '', []
    
    def _extract_guests_from_table(self, table: Table):
        """从表格中提取来宾信息"""
        rows = table.rows
        if len(rows) < 2:
            return
        
        # 检查是否是来宾信息表
        header_text = ' '.join([cell.text for cell in rows[0].cells]).lower()
        if not any(keyword in header_text for keyword in ['来宾', '姓名', '单位', '职务']):
            return
        
        print(f'检测到来宾表格，共{len(rows)}行')
        
        # 从第二行开始提取数据
        for i, row in enumerate(rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            
            # 过滤空行
            if not any(cells):
                continue
            
            print(f'行{i}: {len(cells)}列 -> {cells}')
            
            # 判断第一列是否是序号
            has_seq_no = len(cells) > 0 and re.match(r'^\d+[°º]?$', cells[0])
            start_index = 1 if has_seq_no else 0
            
            guest = None
            
            # 根据列数判断格式
            if len(cells) >= start_index + 5:
                # 6列格式：序号、来宾单位、姓名、民族、职务、健康状况
                guest = GuestInfo(
                    company=cells[start_index] if start_index < len(cells) else '',
                    name=cells[start_index + 1] if start_index + 1 < len(cells) else '',
                    position=cells[start_index + 3] if start_index + 3 < len(cells) else ''  # 跳过民族
                )
                print(f'  提取（6列）: 单位={guest.company}, 姓名={guest.name}, 职务={guest.position}')
                
            elif len(cells) >= start_index + 4:
                # 5列格式：序号、来宾单位、姓名、民族、职务
                guest = GuestInfo(
                    company=cells[start_index] if start_index < len(cells) else '',
                    name=cells[start_index + 1] if start_index + 1 < len(cells) else '',
                    position=cells[start_index + 3] if start_index + 3 < len(cells) else ''
                )
                print(f'  提取（5列）: 单位={guest.company}, 姓名={guest.name}, 职务={guest.position}')
                
            elif len(cells) >= start_index + 3:
                # 4列格式：序号、来宾单位、姓名、职务
                guest = GuestInfo(
                    company=cells[start_index] if start_index < len(cells) else '',
                    name=cells[start_index + 1] if start_index + 1 < len(cells) else '',
                    position=cells[start_index + 2] if start_index + 2 < len(cells) else ''
                )
                print(f'  提取（4列）: 单位={guest.company}, 姓名={guest.name}, 职务={guest.position}')
            
            # 添加来宾信息
            if guest and guest.name and len(guest.name) >= 2:
                # 避免重复
                if not any(g.name == guest.name and g.company == guest.company for g in self.guests):
                    # 确保不是表头
                    if '姓名' not in guest.name and '序号' not in guest.name:
                        self.guests.append(guest)
                        print(f'  ✅ 添加来宾成功')
                    else:
                        print(f'  ⚠️ 跳过（表头）')
                else:
                    print(f'  ⚠️ 跳过（重复）')
        
        print(f'共提取来宾: {len(self.guests)}位')
